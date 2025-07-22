import streamlit as st
import pandas as pd
import os
import yagmail
from fpdf import FPDF
from datetime import datetime
from docx import Document
from docx.shared import Pt
import locale
import re

st.set_page_config(layout="wide", page_title="Gestão de Formulários FRG")

# Configurar locale para formatação de moeda
try:
    locale.setlocale(locale.LC_ALL, 'pt_BR.UTF-8')
except locale.Error:
    st.warning("⚠️ Locale pt_BR.UTF-8 não encontrado. Tentando usar locale padrão.")
    try:
        locale.setlocale(locale.LC_ALL, '')  # Usa o padrão do sistema
    except locale.Error:
        st.error("❌ Nenhum locale adequado encontrado. A formatação de moeda pode não estar correta.")
    

DATA_PATH = "dados_formulario.csv"
EMAIL_REMETENTE = "brunomelo@frg.com.br" # ATUALIZE COM SEU E-MAIL
EMAIL_SENHA = "Trocar@123" # ATUALIZE COM SUA SENHA DE APP DO GMAIL
# --- ETAPAS ATUALIZADAS ---
ETAPAS = ["Aguardando Resposta", "Respondido", "Relação de Crédito", "Desconto de quitação de deficit", "Termo de Portabilidade", "Carta de Portabilidade", "Processo Concluído"]
WORD_TEMPLATE_PATH = "template_quitacao.docx"
WORD_TEMPLATE_PORT_PATH = "template_termo_de_portabilidade.docx"
# --- NOVO TEMPLATE ---
WORD_TEMPLATE_CARTA_PATH = "template_carta.docx" # Certifique-se de que este arquivo existe e é .docx

def carregar_dados():
    colunas_necessarias = [
        "Nome", "Matricula", "CPF", "Email", "Comentário", "Área", "Etapa",
        "Dados Adicionais", "Creditar", "Banco", "Conta", "Agencia", "NomeAgencia",
        "ValorRS", "TipoEntidade", "Patrocinadora", "Plano", "QtdeCotas", "ValorCota",
        "DataValorCota", "MesAnoRelacao", "DataPagamento", "NRefDoc",
        "Rua", "Complemento", "Bairro", "CEP", "Cidade", "UF",
        "MesCalculoCotaDoc", "Deficit2014", "Deficit2022",
        # NOVAS COLUNAS PARA TERMO DE PORTABILIDADE
        "Data_admissao", "Data_desligamento", "Data_inscricao",
        "plano_de_beneficio", "CNPB", "plano_receptor", "cnpj_plano_receptor",
        "endereco_plano_receptor", "cep_plano_receptor", "cidade_plano_receptor",
        "contato_plano_receptor", "telefone_plano_receptor", "email_plano_receptor",
        "banco_plano_receptor", "agencia_plano_receptor", "conta_plano_receptor",
        "Parcela_Participante", "Parcela_Patrocinadora", "Total_acumulado",
        "Regime_de_tributacao", "Recursos_portados", "debito", "total_a_ser_portado",
        "Data_base_portabilidade",
        # --- NOVAS COLUNAS PARA CARTA DE PORTABILIDADE ---
        "Data_de_Transferencia_Carta", "Banco_Carta", "Agencia_Carta", "Conta_Corrente_Carta",
        # --- NOVA COLUNA PARA NUMERO DE RELACAO ---
        "NRelacao" 
    ]
    
    dtype_map = {
        "Nome": str, "Matricula": str, "CPF": str, "Email": str, "Comentário": str, "Área": str, "Etapa": str,
        "Dados Adicionais": str, "Creditar": str, "Banco": str, "Conta": str, "Agencia": str, "NomeAgencia": str,
        "DataValorCota": str, "MesAnoRelacao": str, "DataPagamento": str, "NRefDoc": str,
        "Rua": str, "Complemento": str, "Bairro": str, "CEP": str, "Cidade": str, "UF": str,
        "MesCalculoCotaDoc": str, 
        
        "ValorRS": object, 
        "QtdeCotas": object, 
        "ValorCota": object, 
        "Deficit2014": object, 
        "Deficit2022": object, 
        "Parcela_Participante": object, 
        "Parcela_Patrocinadora": object, 
        "Total_acumulado": object, 
        "Recursos_portados": object, 
        "debito": object, 
        "total_a_ser_portado": object, 
        "NRelacao": object, 

        "Data_admissao": str, "Data_desligamento": str, "Data_inscricao": str,
        "plano_de_beneficio": str, "CNPB": str, "plano_receptor": str, "cnpj_plano_receptor": str,
        "endereco_plano_receptor": str, "cep_plano_receptor": str, "cidade_plano_receptor": str,
        "contato_plano_receptor": str, "telefone_plano_receptor": str, "email_plano_receptor": str,
        "banco_plano_receptor": str, "agencia_plano_receptor": str, "conta_plano_receptor": str,
        "Regime_de_tributacao": str,
        "Data_base_portabilidade": str,
        "Data_de_Transferencia_Carta": str, "Banco_Carta": str, "Agencia_Carta": str, "Conta_Corrente_Carta": str
    }

    if os.path.exists(DATA_PATH):
        try:
            df = pd.read_csv(DATA_PATH, dtype=dtype_map)
            
            for col in colunas_necessarias:
                if col not in df.columns:
                    df[col] = pd.Series(dtype=dtype_map.get(col, object))
            return df
        except pd.errors.EmptyDataError:
            return pd.DataFrame(columns=colunas_necessarias).astype(dtype_map)
        except Exception as e:
            st.error(f"Erro ao carregar o arquivo CSV: {e}")
            return pd.DataFrame(columns=colunas_necessarias).astype(dtype_map)
    else:
        return pd.DataFrame(columns=colunas_necessarias).astype(dtype_map)

def salvar_dados(novo_dado):
    df = carregar_dados()
    novo_dado_df = pd.DataFrame([novo_dado])
    
    for col in df.columns:
        if col not in novo_dado_df.columns:
            novo_dado_df[col] = pd.NA
    
    for col, dtype in df.dtypes.items():
        if col in novo_dado_df.columns and novo_dado_df[col].dtype != dtype:
            try:
                if pd.api.types.is_numeric_dtype(dtype):
                    novo_dado_df[col] = pd.to_numeric(novo_dado_df[col], errors='coerce').astype(dtype)
                else: 
                    novo_dado_df[col] = novo_dado_df[col].astype(dtype)
            except Exception as e:
                st.warning(f"Não foi possível converter a coluna '{col}' para o tipo '{dtype}' ao salvar novos dados: {e}")
                novo_dado_df[col] = pd.NA

    df = pd.concat([df, novo_dado_df[df.columns]], ignore_index=True)
    df.to_csv(DATA_PATH, index=False)


def atualizar_etapa(nome, nova_etapa):
    df = carregar_dados()
    if not df.empty and "Nome" in df.columns:
        df.loc[df["Nome"] == nome, "Etapa"] = nova_etapa
        df.to_csv(DATA_PATH, index=False)

def salvar_dados_completos(nome, dados_dict):
    df = carregar_dados()
    if not df.empty and "Nome" in df.columns:
        idx_list = df[df["Nome"] == nome].index
        if not idx_list.empty:
            idx = idx_list[0]
            for chave, valor in dados_dict.items():
                if chave in df.columns:
                    target_dtype = df[chave].dtype
                    try:
                        if pd.api.types.is_numeric_dtype(target_dtype):
                            df.loc[idx, chave] = pd.to_numeric(valor, errors='coerce').astype(target_dtype)
                        else:
                            df.loc[idx, chave] = str(valor)
                    except Exception as e:
                        st.warning(f"Não foi possível converter a coluna '{chave}' para o tipo '{target_dtype}' ao salvar dados completos: {e}")
                        df.loc[idx, chave] = pd.NA
                else:
                    st.warning(f"Tentativa de salvar coluna inexistente: {chave}")
            df.to_csv(DATA_PATH, index=False)
            return df.loc[idx].to_dict()
    return {}

EMAILS_POR_AREA = {
    "RH": ["rh@empresa.com.br", "gerente.rh@empresa.com.br"],
    "Financeiro": ["financeiro@empresa.com.br", "financeiro.gerencia@empresa.com.br"],
    "Seguridade": ["seguridade@empresa.com.br", "equipe.seguridade@empresa.com.br"],
    "Outra": ["contato.geral@empresa.com.br"]
}

def enviar_email(email_pessoal, nome, destinatarios_internos_selecionados):
    try:
        yag = yagmail.SMTP(EMAIL_REMETENTE, EMAIL_SENHA)

        assunto_participante = f"Confirmação de Recebimento - Formulário FRG - {nome}"
        conteudo_participante = f"Olá {nome},\n\nRecebemos seu formulário com sucesso. Em breve, entraremos em contato ou daremos prosseguimento à sua solicitação.\n\nObrigado,\nEquipe FRG"
        yag.send(to=email_pessoal, subject=assunto_participante, contents=conteudo_participante)
        st.info(f"E-mail de confirmação enviado para **{email_pessoal}**.")

        if not destinatarios_internos_selecionados:
            st.warning(f"⚠️ Nenhum e-mail selecionado para notificação interna. E-mail de notificação não enviado.")
            return
        
        assunto_interno = f"Novo cadastro aguardando resposta - {nome}"
        conteudo_interno = f"Olá,\n\nUm novo formulário foi preenchido por **{nome}** ({email_pessoal}).\n\nPor favor, acesse o sistema para mais detalhes.\n\nAtt,\nSistema Streamlit"
        
        yag.send(to=destinatarios_internos_selecionados, subject=assunto_interno, contents=conteudo_interno)
        st.info(f"E-mail de notificação interna enviado para: **{', '.join(destinatarios_internos_selecionados)}**.")

    except Exception as e:
        st.error(f"❌ Erro ao enviar e-mail: {e}. Verifique as credenciais e as configurações do servidor SMTP.")

def obter_proximo_n_relacao():
    df = carregar_dados()
    if 'NRelacao' in df.columns and pd.to_numeric(df['NRelacao'], errors='coerce').notna().any():
        ultimo_n = pd.to_numeric(df['NRelacao'], errors='coerce').max()
        if pd.isna(ultimo_n): 
            return 1
        return int(ultimo_n) + 1
    return 1

def formatar_matricula(matricula):
    matricula = str(matricula).strip()
    if not matricula:
        return ""
    apenas_digitos = re.sub(r'\D', '', matricula)
    if len(apenas_digitos) > 1:
        return f"{apenas_digitos[:-1]}-{apenas_digitos[-1]}"
    return apenas_digitos

def formatar_conta(numero_conta):
    numero_conta = str(numero_conta).strip()
    if not numero_conta:
        return ""
    apenas_digitos = re.sub(r'\D', '', numero_conta)
    if len(apenas_digitos) > 1 and len(apenas_digitos) <= 12:
        return f"{apenas_digitos[:-1]}-{apenas_digitos[-1]}"
    return apenas_digitos

def formatar_numero_para_exibicao(valor_numerico, casas_decimais=2):
    # # DEPURACAO: # st.write(f"formatar_numero_para_exibicao: Recebido '{valor_numerico}' (tipo: {type(valor_numerico)})")
    try:
        valor_float = float(valor_numerico)
        # # DEPURACAO: # st.write(f"formatar_numero_para_exibicao: Convertido para float: {valor_float}")
        
        formatted_str_us = f"{valor_float:,.{casas_decimais}f}"
        # # DEPURACAO: # st.write(f"formatar_numero_para_exibicao: Formatado US: '{formatted_str_us}'")
        
        final_str = formatted_str_us.replace(",", "X").replace(".", ",").replace("X", ".")
        # # DEPURACAO: # st.write(f"formatar_numero_para_exibicao: Formatado BR Final: '{final_str}'")
        return final_str
    except (ValueError, TypeError) as e:
        # # DEPURACAO: # st.write(f"formatar_numero_para_exibicao: Erro de formatação para '{valor_numerico}': {e}")
        return "0,00"

def formatar_moeda_para_exibicao(valor_numerico):
    # # DEPURACAO: # st.write(f"formatar_moeda_para_exibicao: Recebido '{valor_numerico}' (tipo: {type(valor_numerico)})")
    try:
        float_val = float(valor_numerico)
        # # DEPURACAO: # st.write(f"formatar_moeda_para_exibicao: Convertido para float: {float_val}")
        return locale.currency(float_val, grouping=True, symbol=None)
    except (ValueError, TypeError) as e:
        # # DEPURACAO: # st.write(f"formatar_moeda_para_exibicao: Erro de formatação para '{valor_numerico}': {e}")
        return "0,00"

# --- CRÍTICA: FUNÇÃO desformatar_string_para_float AJUSTADA ---
def desformatar_string_para_float(valor_str):
    # # DEPURACAO: # st.write(f"desformatar_string_para_float: Recebido para desformatação: '{valor_str}' (tipo: {type(valor_str)})")
    if valor_str is None or str(valor_str).strip() == "" or str(valor_str).lower() == 'nan':
        # # DEPURACAO: # st.write(f"desformatar_string_para_float: Retornando 0.0 para vazio/nan.")
        return 0.0
    
    s_val = str(valor_str).strip()

    # Heurística para detectar o formato.
    # Priorizamos o formato BR com vírgula como separador decimal.
    # Se contém vírgula E (não contém ponto OU o ponto está antes da vírgula),
    # assume-se formato BR. Ex: "1.234,56" ou "123,45"
    if ',' in s_val and ('.' not in s_val or s_val.rfind(',') > s_val.rfind('.')):
        # Formato BR: remove pontos de milhar e substitui vírgula decimal por ponto.
        clean_str = s_val.replace('.', '').replace(',', '.')
        # # DEPURACAO: # st.write(f"desformatar_string_para_float: Detectado BR, clean_str: '{clean_str}'")
    else:
        # Assumir formato US (vírgula como milhar, ponto como decimal) ou número puro.
        # Remove apenas vírgulas (separador de milhar US).
        # Ex: "1,234.56" -> "1234.56"
        # Ex: "50" -> "50"
        # Ex: "50.0" -> "50.0" (aqui o ponto é decimal e deve permanecer)
        clean_str = s_val.replace(',', '')
        # # DEPURACAO: # st.write(f"desformatar_string_para_float: Detectado US/Puro, clean_str: '{clean_str}'")

    try:
        float_val = float(clean_str)
        # # DEPURACAO: # st.write(f"desformatar_string_para_float: Converteu para float: {float_val}")
        return float_val
    except ValueError:
        st.warning(f"Não foi possível converter '{valor_str}' para número após desformatação. Usando 0.0.")
        return 0.0

# --- FUNÇÃO PARA SUBSTITUIÇÃO MAIS ROBUSTA (AGORA NO LUGAR CERTO) ---
def replace_placeholders_in_document(doc, substitutions):
    """
    Substitui placeholders em parágrafos e células de tabelas do documento DOCX.
    Esta função tenta ser mais robusta para placeholders que podem estar divididos em runs,
    e tenta preservar o estilo da primeira run.
    """
    def process_paragraph_runs(p, key, value):
        full_text = "".join([run.text for run in p.runs])
        if key in full_text:
            new_full_text = full_text.replace(key, value)
            
            if p.runs:
                first_run_style = p.runs[0].style
                first_run_font = p.runs[0].font
                
                for run in list(p.runs): 
                    p.runs[0]._element.getparent().remove(run._element) 
                
                new_run = p.add_run(new_full_text)
                new_run.style = first_run_style
                new_run.font.name = first_run_font.name
                new_run.font.size = first_run_font.size
                new_run.font.bold = first_run_font.bold
                new_run.font.italic = first_run_font.italic
                new_run.font.underline = first_run_font.underline
            else:
                p.add_run(new_full_text)

    for p in doc.paragraphs:
        for key, value in substitutions.items():
            process_paragraph_runs(p, key, value)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for key, value in substitutions.items():
                        process_paragraph_runs(p, key, value)
    
    for section in doc.sections:
        if section.header:
            for p in section.header.paragraphs:
                for key, value in substitutions.items():
                    process_paragraph_runs(p, key, value)
            for table in section.header.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            for key, value in substitutions.items():
                                process_paragraph_runs(p, key, value)

        if section.footer:
            for p in section.footer.paragraphs:
                for key, value in substitutions.items():
                    process_paragraph_runs(p, key, value)
            for table in section.footer.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            for key, value in substitutions.items():
                                process_paragraph_runs(p, key, value)


def gerar_pdf_relacao_credito(dados):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_font("Arial", 'B', size=12)
    pdf.cell(0, 10, "REAL GRANDEZA", ln=True, align='C')
    pdf.set_font("Arial", 'B', size=10)
    pdf.cell(0, 5, "FUNDAÇÃO DE PREVIDÊNCIA E ASSISTÊNCIA SOCIAL", ln=True, align='C')
    pdf.ln(5)

    mes_ano_relacao = dados.get('MesAnoRelacao', datetime.now().strftime("%b/%y").lower())
    current_y_for_relation = pdf.get_y()
    pdf.set_font("Arial", size=10)

    numero_relacao = dados.get('NRelacao', obter_proximo_n_relacao()) 
    pdf.set_xy(150, current_y_for_relation)
    pdf.multi_cell(50, 5, f"Relação nº {numero_relacao}\n{mes_ano_relacao}", align='R')
    
    pdf.set_xy(10, current_y_for_relation + 5)
    pdf.cell(0, 5, f"GBP/AMX {mes_ano_relacao}", ln=False) 
    pdf.set_y(current_y_for_relation + 10)
    pdf.ln(5)

    pdf.set_font("Arial", '', size=10)
    pdf.cell(0, 7, "DIRETORIA DE SEGURIDADE - DS", ln=True)
    pdf.cell(0, 7, "GERÊNCIA DE BENEFÍCIOS E PAGAMENTOS - GBP", ln=True) 
    pdf.ln(5)
    pdf.set_font("Arial", 'B', size=12)
    pdf.cell(0, 7, "PORTABILIDADE", ln=True, align='C')
    pdf.ln(5)
    pdf.set_font("Arial", size=10)
    
    pdf.cell(30, 7, "Creditar:")
    x_before_cod_banco = pdf.get_x()
    pdf.set_x(120)
    pdf.cell(25, 7, "Cód. do Banco:")
    pdf.set_font("Arial", 'B', size=10)
    pdf.cell(0, 7, str(dados.get('Banco', '')), ln=True)
    pdf.set_x(x_before_cod_banco -20)
    pdf.set_font("Arial", 'B', size=10)
    pdf.cell(0, 7, str(dados.get('Creditar', 'Banco Bradesco')), ln=True)
    
    pdf.set_font("Arial", size=10)
    pdf.cell(15, 7, "Nome:")
    pdf.set_font("Arial", 'B', size=10)
    pdf.cell(0, 7, "Real Grandeza", ln=True)
    
    pdf.set_font("Arial", size=10)
    pdf.cell(0, 7, f"Conta: {formatar_conta(dados.get('Conta', ''))}", ln=True)
    pdf.cell(0, 7, f"Cód. Agência: {dados.get('Agencia', '')}", ln=True)
    pdf.cell(0, 7, f"Nome da Agência: {dados.get('NomeAgencia', '')}", ln=True)
    
    qtde_cotas_val = desformatar_string_para_float(dados.get('QtdeCotas', '0'))
    valor_cota_val = desformatar_string_para_float(dados.get('ValorCota', '0'))
    valor_total_rs_calculado = qtde_cotas_val * valor_cota_val


    pdf.set_font("Arial", 'B', size=10)
    pdf.cell(0, 7, f"Valor em R$: {formatar_moeda_para_exibicao(valor_total_rs_calculado)}", ln=True)
    pdf.set_font("Arial", size=10)
    
    pdf.cell(35, 7, "Tipo de Entidade:"); pdf.set_font("Arial", 'B', size=10); pdf.cell(0, 7, str(dados.get('TipoEntidade', 'Fechada')), ln=True); pdf.set_font("Arial", size=10)
    pdf.cell(35, 7, "PATROCINADORA:"); pdf.set_font("Arial", 'B', size=10); pdf.cell(0, 7, str(dados.get('Patrocinadora', 'FURNAS')), ln=True); pdf.set_font("Arial", size=10)
    pdf.cell(35, 7, "PLANO:"); pdf.set_font("Arial", 'B', size=10); pdf.cell(0, 7, str(dados.get('Plano', 'CONTRIBUIÇÃO DEFINIDA - CD')), ln=True); pdf.set_font("Arial", size=10)
    
    pdf.cell(150, 7, "Total", align='R'); pdf.set_font("Arial", 'B', size=10); pdf.cell(0, 7, formatar_moeda_para_exibicao(valor_total_rs_calculado), ln=True); pdf.set_font("Arial", size=10)
    pdf.ln(3)
    pdf.cell(0, 7, f"Para pagamento dia: {dados.get('DataPagamento', '03/jun/2025')}", ln=True)
    pdf.ln(7)
    pdf.set_font("Arial", 'B', size=11); pdf.cell(0, 7, "Identificação do Participante", ln=True, align='C'); pdf.set_font("Arial", size=10)
    pdf.cell(20, 7, "Nome:"); pdf.set_font("Arial", 'B', size=10); pdf.cell(0, 7, str(dados.get('Nome', '')), ln=True); pdf.set_font("Arial", size=10)
    pdf.cell(20, 7, "Matrícula:"); pdf.set_font("Arial", 'B', size=10); pdf.cell(0, 7, formatar_matricula(dados.get('Matricula', '')), ln=True); pdf.set_font("Arial", size=10)
    pdf.cell(20, 7, "C.P.F.:"); pdf.set_font("Arial", 'B', size=10); pdf.cell(0, 7, str(dados.get('CPF', '')), ln=True); pdf.set_font("Arial", size=10)
    
    pdf.cell(30, 7, "Qtde. de Cotas:"); pdf.set_font("Arial", 'B', size=10); pdf.cell(0, 7, formatar_numero_para_exibicao(qtde_cotas_val, casas_decimais=2), ln=True); pdf.set_font("Arial", size=10)
    data_valor_cota_pdf = dados.get('DataValorCota', '30/04/2025')
    pdf.cell(55, 7, f"Valor da Cota ({data_valor_cota_pdf}):"); pdf.set_font("Arial", 'B', size=10); pdf.cell(0, 7, formatar_numero_para_exibicao(valor_cota_val, casas_decimais=2), ln=True); pdf.set_font("Arial", size=10)

    pdf.ln(10)
    pdf.set_font("Arial", 'I', size=9); pdf.cell(0, 7, "Patrícia Melo e Souza", ln=True, align='C'); pdf.cell(0, 5, "Diretora de Seguridade", ln=True, align='C')
    
    output_filename = f"relacao_credito_{dados.get('Nome', 'Desconhecido').replace(' ', '_')}_N{numero_relacao}.pdf"
    pdf.output(output_filename, 'F')
    return output_filename


def gerar_documento_quitacao(dados_completos):
    if not os.path.exists(WORD_TEMPLATE_PATH):
        st.error(f"Template Word '{WORD_TEMPLATE_PATH}' não encontrado!")
        return None, None
    try:
        doc = Document(WORD_TEMPLATE_PATH)
    except Exception as e:
        st.error(f"Erro ao carregar o template Word '{WORD_TEMPLATE_PATH}': {e}")
        return None, None

    try:
        qtde_cotas = desformatar_string_para_float(dados_completos.get('QtdeCotas', '0'))
        valor_cota = desformatar_string_para_float(dados_completos.get('ValorCota', '0'))
        total_reserva_poupanca_rs = qtde_cotas * valor_cota

        deficit_2014_val = desformatar_string_para_float(dados_completos.get('Deficit2014', '0'))
        deficit_2022_val = desformatar_string_para_float(dados_completos.get('Deficit2022', '0'))
        debito_total_deficit_rs = deficit_2014_val + deficit_2022_val

        anos_com_deficit = []
        if deficit_2014_val > 0:
            anos_com_deficit.append("2014")
        if deficit_2022_val > 0:
            anos_com_deficit.append("2022")
        placeholder_descricao_deficit = " e ".join(anos_com_deficit) if anos_com_deficit else ""

        valor_a_receber_rs = total_reserva_poupanca_rs - debito_total_deficit_rs
    except Exception as e:
        st.error(f"Erro nos cálculos para o documento: {e}. Verifique os valores de cotas e déficits.")
        st.write("Valores usados nos cálculos (após conversão):")
        st.write(f"     Qtde Cotas: {qtde_cotas if 'qtde_cotas' in locals() else 'Erro'}")
        st.write(f"     Valor Cota: {valor_cota if 'valor_cota' in locals() else 'Erro'}")
        st.write(f"     Déficit 2014: {deficit_2014_val if 'deficit_2014_val' in locals() else 'Erro'}")
        st.write(f"     Déficit 2022: {deficit_2022_val if 'deficit_2022_val' in locals() else 'Erro'}")
        return None, None

    substituicoes = {
        "{{N_REF}}": str(dados_completos.get('NRefDoc', '')),
        "{{NOME_PARTICIPANTE}}": str(dados_completos.get('Nome', '')),
        "{{ENDERECO_RUA}}": str(dados_completos.get('Rua', '')),
        "{{ENDERECO_COMPLEMENTO}}": str(dados_completos.get('Complemento', '')).replace('nan', ''),
        "{{ENDERECO_BAIRRO}}": str(dados_completos.get('Bairro', '')),
        "{{ENDERECO_CEP}}": str(dados_completos.get('CEP', '')),
        "{{ENDERECO_CIDADE_UF}}": f"{dados_completos.get('Cidade', '')} - {dados_completos.get('UF', '')}",
        "{{ASSUNTO_MATRICULA}}": formatar_matricula(dados_completos.get('Matricula', '')),
        "{{ASSUNTO_PLANO}}": str(dados_completos.get('Plano', '')),
        "{{ASSUNTO_EMPRESA}}": str(dados_completos.get('Patrocinadora', '')),
        "{{DATA_PAGAMENTO_CREDITO}}": str(dados_completos.get('DataPagamento', '')),
        "{{MES_CALCULO_COTA}}": str(dados_completos.get('MesCalculoCotaDoc', '')),
        "{{SALDO_RESERVA_COTAS}}": formatar_numero_para_exibicao(qtde_cotas, casas_decimais=2),
        "{{VALOR_DA_COTA_RS}}": formatar_numero_para_exibicao(valor_cota, casas_decimais=2),
        "{{TOTAL_RESERVA_POUPANCA_RS}}": formatar_moeda_para_exibicao(total_reserva_poupanca_rs),
        "{{DEBITO_TOTAL_DEFICIT_RS}}": formatar_moeda_para_exibicao(debito_total_deficit_rs),
        "{{DESCRICAO_DEFICIT}}": placeholder_descricao_deficit,
        "{{VALOR_A_RECEBER_RS}}": formatar_moeda_para_exibicao(valor_a_receber_rs)
    }

    st.write("---")
    st.write("### Debugging: Dicionário de Substituições para Quitação")
    st.json(substituicoes)
    st.write("---")

    replace_placeholders_in_document(doc, substituicoes)

    nome_base = str(dados_completos.get('Nome', 'Desconhecido')).replace(' ', '_').replace('/', '_')
    output_docx_path = f"quitacao_deficit_{nome_base}.docx"
    output_pdf_path = f"quitacao_deficit_{nome_base}.pdf"

    try:
        doc.save(output_docx_path)
        st.info(f"Arquivo DOCX '{output_docx_path}' gerado.")
        st.warning("A conversão para PDF pode não funcionar em ambientes de nuvem (Streamlit Cloud) sem o Microsoft Word ou LibreOffice instalado.")
        return None, output_docx_path
    except Exception as e_docx:
        st.error(f"Erro ao salvar o documento DOCX: {e_docx}")
        return None, None

def gerar_documento_portabilidade(dados_completos):
    if not os.path.exists(WORD_TEMPLATE_PORT_PATH):
        st.error(f"Template Word de Portabilidade '{WORD_TEMPLATE_PORT_PATH}' não encontrado!")
        return None, None
    try:
        doc = Document(WORD_TEMPLATE_PORT_PATH)
    except Exception as e:
        st.error(f"Erro ao carregar o template Word de Portabilidade '{WORD_TEMPLATE_PORT_PATH}': {e}")
        return None, None

    try:
        parcela_participante_val = desformatar_string_para_float(dados_completos.get('Parcela_Participante', '0'))
        parcela_patrocinadora_val = desformatar_string_para_float(dados_completos.get('Parcela_Patrocinadora', '0'))
        debito_val = desformatar_string_para_float(dados_completos.get('debito', '0'))

        total_acumulado_val = parcela_participante_val + parcela_patrocinadora_val
        valor_total_a_ser_portado_val = total_acumulado_val - debito_val

    except Exception as e:
        st.error(f"Erro nos cálculos para o Termo de Portabilidade: {e}. Verifique os valores monetários.")
        return None, None
    
    substituicoes = {
        "{{NOME_PARTICIPANTE}}": str(dados_completos.get('Nome', '')),
        "{{CPF}}": str(dados_completos.get('CPF', '')),
        "{{Matricula}}": formatar_matricula(dados_completos.get('Matricula', '')),
        "{{ENDERECO_COMPLEMENTO}}": str(dados_completos.get('Complemento', '')).replace('nan', ''),
        "{{ENDERECO_RUA}}": str(dados_completos.get('Rua', '')),
        "{{ENDERECO_BAIRRO}}": str(dados_completos.get('Bairro', '')),
        "{{ENDERECO_CEP}}": str(dados_completos.get('CEP', '')),
        "{{ENDERECO_CIDADE_UF}}": f"{dados_completos.get('Cidade', '')} - {dados_completos.get('UF', '')}",
        "{{ASSUNTO_EMPRESA}}": str(dados_completos.get('Patrocinadora', '')),
        "{{Data_admissao}}": str(dados_completos.get('Data_admissao', '')),
        "{{Data_desligamento}}": str(dados_completos.get('Data_desligamento', '')),
        "{{Data_inscricao}}": str(dados_completos.get('Data_inscricao', '')),
        "{{plano_de_beneficio}}": str(dados_completos.get('plano_de_beneficio', '')),
        "{{CNPB}}": str(dados_completos.get('CNPB', '')),
        "{{plano_receptor}}": str(dados_completos.get('plano_receptor', '')),
        "{{cnpj_plano_receptor}}": str(dados_completos.get('cnpj_plano_receptor', '')),
        "{{endereco_plano_receptor}}": str(dados_completos.get('endereco_plano_receptor', '')),
        "{{cep_plano_receptor}}": str(dados_completos.get('cep_plano_receptor', '')),
        "{{cidade_plano_receptor}}": str(dados_completos.get('cidade_plano_receptor', '')),
        "{{contato_plano_receptor}}": str(dados_completos.get('contato_plano_receptor', '')),
        "{{telefone_plano_receptor}}": str(dados_completos.get('telefone_plano_receptor', '')),
        "{{email_plano_receptor}}": str(dados_completos.get('email_plano_receptor', '')),
        "{{banco_plano_receptor}}": str(dados_completos.get('banco_plano_receptor', '')),
        "{{agencia_plano_receptor}}": str(dados_completos.get("agencia_plano_receptor", '')),
        "{{conta_plano_receptor}}": str(dados_completos.get("conta_plano_receptor", '')),
        "{{Parcela_Participante}}": formatar_moeda_para_exibicao(parcela_participante_val),
        "{{Parcela_Patrocinadora}}": formatar_moeda_para_exibicao(parcela_patrocinadora_val),
        "{{Total_acumulado}}": formatar_moeda_para_exibicao(total_acumulado_val),
        "{{Regime_de_tributacao}}": str(dados_completos.get('Regime_de_tributacao', '')),
        "{{Recursos_portados}}": formatar_moeda_para_exibicao(desformatar_string_para_float(dados_completos.get('Recursos_portados', '0'))),
        "{{debito}}": formatar_moeda_para_exibicao(debito_val),
        "{{total_a_ser_portado}}": formatar_moeda_para_exibicao(valor_total_a_ser_portado_val),
        "{{Data_base}}": str(dados_completos.get('Data_base_portabilidade', ''))
    }

    st.write("---")
    st.write("### Debugging: Dicionário de Substituições para Termo de Portabilidade")
    st.json(substituicoes)
    st.write("---")

    replace_placeholders_in_document(doc, substituicoes)

    nome_base = str(dados_completos.get('Nome', 'Desconhecido')).replace(' ', '_').replace('/', '_')
    output_docx_path = f"termo_portabilidade_{nome_base}.docx"
    output_pdf_path = f"termo_portabilidade_{nome_base}.pdf"

    try:
        doc.save(output_docx_path)
        st.info(f"Arquivo DOCX '{output_docx_path}' gerado.")
        st.warning("A conversão para PDF pode não funcionar em ambientes de nuvem (Streamlit Cloud) sem o Microsoft Word ou LibreOffice instalado.")
        return None, output_docx_path
    except Exception as e_docx:
        st.error(f"Erro ao salvar o documento DOCX: {e_docx}")
        return None, None

def gerar_documento_carta_portabilidade(dados_completos):
    if not os.path.exists(WORD_TEMPLATE_CARTA_PATH):
        st.error(f"Template Word da Carta de Portabilidade '{WORD_TEMPLATE_CARTA_PATH}' não encontrado! Por favor, converta seu template .doc para .docx.")
        return None, None
    try:
        doc = Document(WORD_TEMPLATE_CARTA_PATH)
    except Exception as e:
        st.error(f"Erro ao carregar o template Word da Carta de Portabilidade '{WORD_TEMPLATE_CARTA_PATH}': {e}")
        return None, None

    data_transferencia = str(dados_completos.get('Data_de_Transferencia_Carta', ''))
    banco_carta = str(dados_completos.get('Banco_Carta', ''))
    agencia_carta = str(dados_completos.get('Agencia_Carta', ''))
    conta_corrente_carta = str(dados_completos.get('Conta_Corrente_Carta', ''))

    substituicoes = {
        "{{NOME_PARTICIPANTE}}": str(dados_completos.get('Nome', '')),
        # CORREÇÃO AQUI: Garante que o valor é string antes de chamar replace
        "{{ENDERECO_COMPLEMENTO}}": str(dados_completos.get('Complemento', '')).replace('nan', ''), 
        "{{ENDERECO_RUA}}": str(dados_completos.get('Rua', '')),
        "{{ENDERECO_BAIRRO}}": str(dados_completos.get('Bairro', '')),
        "{{ENDERECO_CEP}}": str(dados_completos.get('CEP', '')),
        "{{ENDERECO_CIDADE_UF}}": f"{dados_completos.get('Cidade', '')} - {dados_completos.get('UF', '')}",
        "{{ASSUNTO_PLANO}}": str(dados_completos.get('Plano', '')),
        "{{ASSUNTO_MATRICULA}}": formatar_matricula(dados_completos.get('Matricula', '')),
        "{{CPF}}": str(dados_completos.get('CPF', '')),
        "{{DATA_DE_TRANSFERENCIA}}": data_transferencia,
        "{{BANCO}}": banco_carta,
        "{{AGENCIA}}": agencia_carta,
        "{{CONTA_CORRENTE}}": formatar_conta(conta_corrente_carta),
        "{{N_Ref}}": str(dados_completos.get('NRefDoc', '')),
        "{{DATA_ATUAL_CARTA}}": datetime.now().strftime("%d de %B de %Y").replace(
            'January', 'janeiro').replace('February', 'fevereiro').replace('March', 'março').replace(
            'April', 'abril').replace('May', 'maio').replace('June', 'junho').replace(
            'July', 'julho').replace('August', 'agosto').replace('September', 'setembro').replace(
            'October', 'outubro').replace('November', 'novembro').replace('December', 'dezembro'),
    }
    
    st.write("---")
    st.write("### Debugging: Dicionário de Substituições para Carta de Portabilidade")
    st.json(substituicoes)
    st.write("---")

    replace_placeholders_in_document(doc, substituicoes)

    nome_base = str(dados_completos.get('Nome', 'Desconhecido')).replace(' ', '_').replace('/', '_')
    output_docx_path = f"carta_portabilidade_{nome_base}.docx"
    output_pdf_path = f"carta_portabilidade_{nome_base}.pdf"

    try:
        doc.save(output_docx_path)
        st.info(f"Arquivo DOCX '{output_docx_path}' gerado.")
        st.warning("A conversão para PDF pode não funcionar em ambientes de nuvem (Streamlit Cloud) sem o Microsoft Word ou LibreOffice instalado.")
        return None, output_docx_path
    except Exception as e_docx:
        st.error(f"Erro ao salvar o documento DOCX: {e_docx}")
        return None, None


if 'download_pdf_relacao' not in st.session_state:
    st.session_state.download_pdf_relacao = None
if 'download_docx_quitacao' not in st.session_state:
    st.session_state.download_docx_quitacao = None
if 'download_docx_portabilidade' not in st.session_state:
    st.session_state.download_docx_portabilidade = None
if 'download_docx_carta' not in st.session_state:
    st.session_state.download_docx_carta = None

tab1, tab2, tab3, tab4, tab5, tab6 = st.tabs(["📥 Formulário Inicial", "📊 Kanban", "📝 Relação de Crédito", "📉 Desconto de Déficit", "📄 Termo de Portabilidade", "📧 Carta de Portabilidade"])


with tab1:
    st.header("📥 Preencha o Formulário Inicial")
    with st.form("form_inicial_tab1"):
        nome_t1 = st.text_input("Nome Completo", key="nome_t1")
        matricula_t1 = st.text_input("Matrícula", key="mat_t1")
        cpf_t1 = st.text_input("CPF", key="cpf_t1")
        email_t1 = st.text_input("Email Contato", key="email_t1")
        comentario_t1 = st.text_area("Comentário", key="com_t1")
        
        area_t1 = st.selectbox("Área Principal do Formulário", list(EMAILS_POR_AREA.keys()), key="area_t1_sb")
        
        todos_os_emails_internos = []
        for emails_lista in EMAILS_POR_AREA.values():
            todos_os_emails_internos.extend(emails_lista)
        todos_os_emails_internos.extend(["contato.geral.adicional@empresa.com.br", "outro.diretor@empresa.com.br"]) 
        todos_os_emails_internos = sorted(list(set(todos_os_emails_internos)))

        default_emails_selecionados = st.session_state.get(f"area_t1_default_emails_{area_t1}", EMAILS_POR_AREA.get(area_t1, []))

        destinatarios_internos_selecionados = st.multiselect(
            "Destinatários Internos da Notificação (inclui e-mails da área selecionada por padrão)",
            options=todos_os_emails_internos,
            default=default_emails_selecionados,
            help="Escolha outros endereços de e-mail que devem receber a notificação sobre este formulário, além dos e-mails da área principal selecionada.",
            key=f"dest_internos_multi_{area_t1}"
        )
        
        st.session_state[f"area_t1_default_emails_{area_t1}"] = destinatarios_internos_selecionados

        enviado_t1 = st.form_submit_button("🚀 Enviar")
        if enviado_t1:
            if nome_t1 and email_t1 and cpf_t1:
                novo_dado = {col: pd.NA for col in carregar_dados().columns}
                novo_dado.update({
                    "Nome": nome_t1, "Matricula": matricula_t1, "CPF": cpf_t1, "Email": email_t1,
                    "Comentário": comentario_t1, "Área": area_t1, "Etapa": "Aguardando Resposta",
                    "MesAnoRelacao": datetime.now().strftime("%b/%y").lower(),
                    "Deficit2014": "0,00", "Deficit2022": "0,00",
                    "Parcela_Participante": "0,00", "Parcela_Patrocinadora": "0,00",
                    "Total_acumulado": "0,00", "Recursos_portados": "0,00", "debito": "0,00",
                    "total_a_ser_portado": "0,00",
                    "NRelacao": obter_proximo_n_relacao()
                })
                salvar_dados(novo_dado)
                st.success(f"✅ Dados de {nome_t1} salvos!")
                
                enviar_email(email_t1, nome_t1, destinatarios_internos_selecionados) 
                st.rerun()
            else:
                st.warning("⚠️ Preencha Nome, CPF e Email.")

with tab2: # KANBAN
    st.header("📌 Painel Kanban")
    
    etapas_para_kanban = [e for e in ETAPAS if e != "Processo Concluído"]
    colunas_kanban = st.columns(len(etapas_para_kanban))
    df_kanban = carregar_dados()

    for i, etapa_k in enumerate(etapas_para_kanban):
        with colunas_kanban[i]:
            etapa_df_k = df_kanban[df_kanban["Etapa"] == etapa_k] if "Etapa" in df_kanban.columns else pd.DataFrame()
            st.subheader(f"{etapa_k} ({len(etapa_df_k)})")
            
            etapa_df_k = etapa_df_k.sort_values(by="Nome", ascending=True)

            for idx_k, row_k in etapa_df_k.iterrows():
                key_base_k = f"{row_k.get('Nome','key')}_{idx_k}_{etapa_k.replace(' ','_')}"
                with st.expander(f"{row_k.get('Nome','Sem Nome')} ({row_k.get('Area','N/A')})", expanded=False):
                    st.caption(f"Matrícula: {formatar_matricula(row_k.get('Matricula', 'N/A'))} | CPF: {row_k.get('CPF', 'N/A')}")
                    
                    if etapa_k == "Aguardando Resposta":
                        if st.button("✅ Respondido", key=f"resp_{key_base_k}"):
                            atualizar_etapa(row_k["Nome"], "Respondido"); st.rerun()
                    elif etapa_k == "Respondido":
                        if st.button("➡️ Relação Crédito", key=f"rel_{key_base_k}"):
                            atualizar_etapa(row_k["Nome"], "Relação de Crédito"); st.rerun()
                        if st.button("➡️ Termo Portabilidade", key=f"port_k_{key_base_k}"):
                            atualizar_etapa(row_k["Nome"], "Termo de Portabilidade"); st.rerun()
                        if st.button("➡️ Carta de Portabilidade", key=f"carta_k_{key_base_k}"): 
                            atualizar_etapa(row_k["Nome"], "Carta de Portabilidade"); st.rerun()
                    elif etapa_k == "Relação de Crédito":
                        if st.button("➡️ Desconto Déficit", key=f"desc_{key_base_k}"):
                            atualizar_etapa(row_k["Nome"], "Desconto de quitação de deficit"); st.rerun()
                        if st.button("⏪ Voltar para Respondido", key=f"volt_resp_rel_{key_base_k}"):
                            atualizar_etapa(row_k["Nome"], "Respondido"); st.rerun()
                    elif etapa_k == "Desconto de quitação de deficit":
                        st.info("Preencher na Aba 'Desconto de Déficit'")
                        if st.button("➡️ Termo Portabilidade", key=f"next_port_from_desc_{key_base_k}"):
                            atualizar_etapa(row_k["Nome"], "Termo de Portabilidade"); st.rerun()
                        if st.button("⏪ Voltar para Relação de Crédito", key=f"volt_rel_desc_{key_base_k}"):
                            atualizar_etapa(row_k["Nome"], "Relação de Crédito"); st.rerun()
                    elif etapa_k == "Termo de Portabilidade":
                        st.info("Preencher na Aba 'Termo de Portabilidade'")
                        if st.button("➡️ Carta de Portabilidade", key=f"next_carta_from_termo_{key_base_k}"):
                            atualizar_etapa(row_k["Nome"], "Carta de Portabilidade"); st.rerun()
                        if st.button("⏪ Voltar para Desconto de Quitação", key=f"volt_desc_termo_{key_base_k}"):
                            atualizar_etapa(row_k["Nome"], "Desconto de quitação de deficit"); st.rerun()
                    elif etapa_k == "Carta de Portabilidade":
                        st.info("Preencher na Aba 'Carta de Portabilidade'")
                        if st.button("✔️ Concluído", key=f"concluido_{key_base_k}"):
                            atualizar_etapa(row_k["Nome"], "Processo Concluído")
                            st.success(f"Processo para {row_k['Nome']} marcado como concluído!")
                            st.rerun()
                        if st.button("⏪ Voltar para Termo de Portabilidade", key=f"volt_termo_carta_{key_base_k}"):
                            atualizar_etapa(row_k["Nome"], "Termo de Portabilidade"); st.rerun()


with tab3: # Relação de Crédito
    st.header("📝 Relação de Crédito")
    df_relacao = carregar_dados()
    df_relacao_credito = df_relacao[df_relacao["Etapa"] == "Relação de Crédito"]

    if df_relacao_credito.empty:
        st.info("Nenhum formulário na etapa 'Relação de Crédito'.")
    else:
        for idx, row in df_relacao_credito.iterrows():
            with st.expander(f"Detalhes de {row['Nome']} - Matrícula: {formatar_matricula(row['Matricula'])}", expanded=False):
                with st.form(f"form_relacao_credito_{row['Nome']}"):
                    st.write(f"Preenchendo dados para **{row['Nome']}** (Matrícula: {formatar_matricula(row['Matricula'])})")

                    creditar = st.text_input("Creditar", value=row.get('Creditar', 'Banco Bradesco'), key=f"cred_{idx}")
                    banco = st.text_input("Banco (Código)", value=row.get('Banco', ''), key=f"bank_{idx}")
                    conta = st.text_input("Conta", value=row.get('Conta', ''), key=f"acc_{idx}")
                    agencia = st.text_input("Agência (Código)", value=row.get('Agencia', ''), key=f"ag_{idx}")
                    nome_agencia = st.text_input("Nome da Agência", value=row.get('NomeAgencia', ''), key=f"nag_{idx}")
                    
                    tipo_entidade = st.text_input("Tipo de Entidade", value=row.get('TipoEntidade', 'Fechada'), key=f"te_{idx}")
                    patrocinadora = st.text_input("Patrocinadora", value=row.get('Patrocinadora', 'FURNAS'), key=f"pat_{idx}")
                    plano = st.text_input("Plano", value=row.get('Plano', 'CONTRIBUIÇÃO DEFINIDA - CD'), key=f"plano_{idx}")
                    
                    qtde_cotas_str = st.text_input("Qtde. de Cotas", value=str(desformatar_string_para_float(row.get('QtdeCotas', '0'))), key=f"qtde_cotas_{idx}")
                    valor_cota_str = st.text_input("Valor da Cota (R$)", value=str(desformatar_string_para_float(row.get('ValorCota', '0'))), key=f"valor_cota_{idx}")

                    data_valor_cota = st.text_input("Data do Valor da Cota (dd/mm/aaaa)", value=row.get('DataValorCota', ''), key=f"data_vc_{idx}")

                    mes_ano_relacao = st.text_input("Mês/Ano da Relação (ex: jun/25)", value=row.get('MesAnoRelacao', datetime.now().strftime("%b/%y").lower()), key=f"mar_{idx}")
                    data_pagamento = st.text_input("Data de Pagamento (dd/mm/aaaa)", value=row.get('DataPagamento', ''), key=f"dp_{idx}")
                    
                    n_ref_doc = st.text_input("Nº Ref. Documento (para Quitação)", value=row.get('NRefDoc', ''), key=f"nref_{idx}")


                    col_rua, col_comp = st.columns(2)
                    rua = col_rua.text_input("Rua", value=row.get('Rua', ''), key=f"rua_{idx}")
                    # CORREÇÃO AQUI: Garante que o valor é string antes de chamar replace
                    complemento = col_comp.text_input("Complemento", value=str(row.get('Complemento', '')).replace('nan', ''), key=f"comp_{idx}")
                    
                    col_bairro, col_cep = st.columns(2)
                    bairro = col_bairro.text_input("Bairro", value=row.get('Bairro', ''), key=f"bairro_{idx}")
                    cep = col_cep.text_input("CEP", value=row.get('CEP', ''), key=f"cep_{idx}")
                    
                    col_cidade, col_uf = st.columns(2)
                    cidade = col_cidade.text_input("Cidade", value=row.get('Cidade', ''), key=f"cidade_{idx}")
                    uf = col_uf.text_input("UF", value=row.get('UF', ''), key=f"uf_{idx}")

                    mes_calculo_cota_doc = st.text_input("Mês de Cálculo Cota (Doc)", value=row.get('MesCalculoCotaDoc', ''), key=f"mes_calc_cota_doc_{idx}")

                    submitted_relacao = st.form_submit_button("💾 Salvar Dados e Gerar Documento")
                    if submitted_relacao:
                        dados_atualizados = {
                            "Creditar": creditar, "Banco": banco, "Conta": conta, "Agencia": agencia, "NomeAgencia": nome_agencia,
                            "TipoEntidade": tipo_entidade, "Patrocinadora": patrocinadora, "Plano": plano,
                            "QtdeCotas": desformatar_string_para_float(qtde_cotas_str), 
                            "ValorCota": desformatar_string_para_float(valor_cota_str), 
                            "DataValorCota": data_valor_cota, "MesAnoRelacao": mes_ano_relacao, "DataPagamento": data_pagamento,
                            "NRefDoc": n_ref_doc,
                            "Rua": rua, "Complemento": complemento, "Bairro": bairro, "CEP": cep, "Cidade": cidade, "UF": uf,
                            "MesCalculoCotaDoc": mes_calculo_cota_doc
                        }
                        
                        salvar_dados_completos(row['Nome'], dados_atualizados)
                        st.success(f"Dados de Relação de Crédito para {row['Nome']} salvos!")
                        
                        pdf_path = gerar_pdf_relacao_credito(row.to_dict())
                        st.session_state.download_pdf_relacao = pdf_path
                        st.rerun()
                
                if st.session_state.download_pdf_relacao:
                    if os.path.exists(st.session_state.download_pdf_relacao):
                        with open(st.session_state.download_pdf_relacao, "rb") as file:
                            st.download_button(
                                label="📥 Download Relação de Crédito PDF",
                                data=file,
                                file_name=os.path.basename(st.session_state.download_pdf_relacao),
                                mime="application/pdf",
                                key=f"download_btn_relacao_{idx}"
                            )
                    else:
                        st.error(f"Arquivo PDF não encontrado em: {st.session_state.download_pdf_relacao}")

with tab4: # Desconto de Quitação de Déficit
    st.header("📉 Desconto de Quitação de Déficit")
    df_deficit = carregar_dados()
    df_deficit_quitacao = df_deficit[df_deficit["Etapa"] == "Desconto de quitação de deficit"]

    if df_deficit_quitacao.empty:
        st.info("Nenhum formulário na etapa 'Desconto de quitação de deficit'.")
    else:
        for idx, row in df_deficit_quitacao.iterrows():
            with st.expander(f"Detalhes de {row['Nome']} - Matrícula: {formatar_matricula(row['Matricula'])}", expanded=False):
                with st.form(f"form_deficit_{row['Nome']}"):
                    st.write(f"Preenchendo dados de quitação de déficit para **{row['Nome']}** (Matrícula: {formatar_matricula(row['Matricula'])})")

                    deficit_2014_str = st.text_input("Déficit 2014 (R$)", value=formatar_moeda_para_exibicao(desformatar_string_para_float(row.get('Deficit2014', '0'))), key=f"def14_{idx}")
                    deficit_2022_str = st.text_input("Déficit 2022 (R$)", value=formatar_moeda_para_exibicao(desformatar_string_para_float(row.get('Deficit2022', '0'))), key=f"def22_{idx}")
                    
                    n_ref_doc = st.text_input("Nº Ref. Documento", value=row.get('NRefDoc', ''), key=f"nref_q_{idx}")
                    data_pagamento = st.text_input("Data de Pagamento (dd/mm/aaaa)", value=row.get('DataPagamento', ''), key=f"dp_q_{idx}")
                    mes_calculo_cota_doc = st.text_input("Mês de Cálculo Cota (Doc)", value=row.get('MesCalculoCotaDoc', ''), key=f"mes_calc_cota_q_doc_{idx}")
                    
                    qtde_cotas_str = st.text_input("Qtde. de Cotas", value=str(desformatar_string_para_float(row.get('QtdeCotas', '0'))), key=f"qtde_cotas_q_{idx}")
                    valor_cota_str = st.text_input("Valor da Cota (R$)", value=str(desformatar_string_para_float(row.get('ValorCota', '0'))), key=f"valor_cota_q_{idx}")
                    data_valor_cota = st.text_input("Data do Valor da Cota (dd/mm/aaaa)", value=row.get('DataValorCota', ''), key=f"data_vc_q_{idx}")


                    submitted_deficit = st.form_submit_button("💾 Salvar Dados e Gerar DOCX (Quitação)")
                    if submitted_deficit:
                        dados_atualizados = {
                            "Deficit2014": desformatar_string_para_float(deficit_2014_str), 
                            "Deficit2022": desformatar_string_para_float(deficit_2022_str), 
                            "NRefDoc": n_ref_doc,
                            "DataPagamento": data_pagamento,
                            "MesCalculoCotaDoc": mes_calculo_cota_doc,
                            "QtdeCotas": desformatar_string_para_float(qtde_cotas_str),
                            "ValorCota": desformatar_string_para_float(valor_cota_str),
                            "DataValorCota": data_valor_cota
                        }
                        dados_completos_apos_salvar = salvar_dados_completos(row['Nome'], dados_atualizados)
                        st.success(f"Dados de Desconto de Quitação para {row['Nome']} salvos!")
                        
                        pdf_path_quit, docx_path_quit = gerar_documento_quitacao(dados_completos_apos_salvar)
                        st.session_state.download_docx_quitacao = docx_path_quit
                        st.rerun()
                
                if st.session_state.download_docx_quitacao:
                    if os.path.exists(st.session_state.download_docx_quitacao):
                        with open(st.session_state.download_docx_quitacao, "rb") as file:
                            st.download_button(
                                label="📥 Download Quitação DOCX",
                                data=file,
                                file_name=os.path.basename(st.session_state.download_docx_quitacao),
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                key=f"download_btn_quitacao_{idx}"
                            )
                    else:
                        st.error(f"Arquivo DOCX não encontrado em: {st.session_state.download_docx_quitacao}")

with tab5: # Termo de Portabilidade
    st.header("📄 Termo de Portabilidade")
    df_portabilidade = carregar_dados()
    df_termo_portabilidade = df_portabilidade[df_portabilidade["Etapa"] == "Termo de Portabilidade"]

    if df_termo_portabilidade.empty:
        st.info("Nenhum formulário na etapa 'Termo de Portabilidade'.")
    else:
        for idx, row in df_termo_portabilidade.iterrows():
            with st.expander(f"Detalhes de {row['Nome']} - Matrícula: {formatar_matricula(row['Matricula'])}", expanded=False):
                with st.form(f"form_termo_portabilidade_{row['Nome']}"):
                    st.write(f"Preenchendo dados para Termo de Portabilidade para **{row['Nome']}** (Matrícula: {formatar_matricula(row['Matricula'])})")

                    st.subheader("Dados do Participante (revisar)")
                    st.text_input("Nome", value=row.get('Nome', ''), disabled=True)
                    st.text_input("CPF", value=row.get('CPF', ''), disabled=True)
                    st.text_input("Matrícula", value=formatar_matricula(row.get('Matricula', '')), disabled=True)
                    st.text_input("Rua", value=row.get('Rua', ''), key=f"rua_port_{idx}")
                    # CORREÇÃO AQUI: Garante que o valor é string antes de chamar replace
                    complemento = st.text_input("Complemento", value=str(row.get('Complemento', '')).replace('nan', ''), key=f"comp_port_{idx}")
                    st.text_input("Bairro", value=row.get('Bairro', ''), key=f"bairro_port_{idx}")
                    st.text_input("CEP", value=row.get('CEP', ''), key=f"cep_port_{idx}")
                    st.text_input("Cidade", value=row.get('Cidade', ''), key=f"cidade_port_{idx}")
                    st.text_input("UF", value=row.get('UF', ''), key=f"uf_port_{idx}")

                    st.subheader("Dados de Admissão/Desligamento/Inscrição")
                    data_admissao = st.text_input("Data de Admissão (dd/mm/aaaa)", value=row.get('Data_admissao', ''), key=f"data_adm_{idx}")
                    data_desligamento = st.text_input("Data de Desligamento (dd/mm/aaaa)", value=row.get('Data_desligamento', ''), key=f"data_desl_{idx}")
                    data_inscricao = st.text_input("Data de Inscrição no Plano (dd/mm/aaaa)", value=row.get('Data_inscricao', ''), key=f"data_insc_{idx}")

                    st.subheader("Dados do Plano de Benefício e Receptor")
                    plano_de_beneficio = st.text_input("Plano de Benefício (Origem)", value=row.get('plano_de_beneficio', ''), key=f"plano_ben_{idx}")
                    cnpb = st.text_input("CNPB (Plano Origem)", value=row.get('CNPB', ''), key=f"cnpb_{idx}")
                    plano_receptor = st.text_input("Plano Receptor (Destino)", value=row.get('plano_receptor', ''), key=f"plano_rec_{idx}")
                    cnpj_plano_receptor = st.text_input("CNPJ do Plano Receptor", value=row.get('cnpj_plano_receptor', ''), key=f"cnpj_rec_{idx}")
                    endereco_plano_receptor = st.text_input("Endereço do Plano Receptor", value=row.get('endereco_plano_receptor', ''), key=f"end_rec_{idx}")
                    cep_plano_receptor = st.text_input("CEP do Plano Receptor", value=row.get('cep_plano_receptor', ''), key=f"cep_rec_{idx}")
                    cidade_plano_receptor = st.text_input("Cidade do Plano Receptor", value=row.get('cidade_plano_receptor', ''), key=f"cidade_rec_{idx}")
                    contato_plano_receptor = st.text_input("Contato do Plano Receptor", value=row.get('contato_plano_receptor', ''), key=f"cont_rec_{idx}")
                    telefone_plano_receptor = st.text_input("Telefone do Plano Receptor", value=row.get('telefone_plano_receptor', ''), key=f"tel_rec_{idx}")
                    email_plano_receptor = st.text_input("Email do Plano Receptor", value=row.get('email_plano_receptor', ''), key=f"email_rec_{idx}")
                    banco_plano_receptor = st.text_input("Banco do Plano Receptor", value=row.get('banco_plano_receptor', ''), key=f"banco_rec_{idx}")
                    agencia_plano_receptor = st.text_input("Agência do Plano Receptor", value=row.get('agencia_plano_receptor', ''), key=f"ag_rec_{idx}")
                    conta_plano_receptor = st.text_input("Conta do Plano Receptor", value=row.get('conta_plano_receptor', ''), key=f"conta_rec_{idx}")

                    st.subheader("Valores e Tributação")
                    parcela_participante_str = st.text_input("Parcela do Participante (R$)", value=formatar_moeda_para_exibicao(desformatar_string_para_float(row.get('Parcela_Participante', '0'))), key=f"par_part_{idx}")
                    parcela_patrocinadora_str = st.text_input("Parcela da Patrocinadora (R$)", value=formatar_moeda_para_exibicao(desformatar_string_para_float(row.get('Parcela_Patrocinadora', '0'))), key=f"par_pat_{idx}")
                    regime_tributacao = st.selectbox("Regime de Tributação", ["Regressivo", "Progressivo", "Não Definido"], index=["Regressivo", "Progressivo", "Não Definido"].index(row.get('Regime_de_tributacao', 'Não Definido')) if row.get('Regime_de_tributacao', 'Não Definido') in ["Regressivo", "Progressivo", "Não Definido"] else 2, key=f"reg_trib_{idx}")
                    recursos_portados_str = st.text_input("Recursos a serem portados (R$)", value=formatar_moeda_para_exibicao(desformatar_string_para_float(row.get('Recursos_portados', '0'))), key=f"rec_port_{idx}")
                    debito_str = st.text_input("Débito (R$)", value=formatar_moeda_para_exibicao(desformatar_string_para_float(row.get('debito', '0'))), key=f"debito_port_{idx}")
                    data_base_portabilidade = st.text_input("Data Base Portabilidade (dd/mm/aaaa)", value=row.get('Data_base_portabilidade', ''), key=f"data_base_port_{idx}")

                    submitted_termo_port = st.form_submit_button("💾 Salvar Dados e Gerar DOCX (Termo de Portabilidade)")
                    if submitted_termo_port:
                        dados_atualizados = {
                            "Rua": st.session_state[f"rua_port_{idx}"], "Complemento": st.session_state[f"comp_port_{idx}"],
                            "Bairro": st.session_state[f"bairro_port_{idx}"], "CEP": st.session_state[f"cep_port_{idx}"],
                            "Cidade": st.session_state[f"cidade_port_{idx}"], "UF": st.session_state[f"uf_port_{idx}"],
                            "Data_admissao": data_admissao, "Data_desligamento": data_desligamento, "Data_inscricao": data_inscricao,
                            "plano_de_beneficio": plano_de_beneficio, "CNPB": cnpb, "plano_receptor": plano_receptor, "cnpj_plano_receptor": cnpj_plano_receptor,
                            "endereco_plano_receptor": endereco_plano_receptor, "cep_plano_receptor": cep_plano_receptor, "cidade_plano_receptor": cidade_plano_receptor,
                            "contato_plano_receptor": contato_plano_receptor, "telefone_plano_receptor": telefone_plano_receptor, "email_plano_receptor": email_plano_receptor,
                            "banco_plano_receptor": banco_plano_receptor, "agencia_plano_receptor": agencia_plano_receptor, "conta_plano_receptor": conta_plano_receptor,
                            "Parcela_Participante": desformatar_string_para_float(parcela_participante_str),
                            "Parcela_Patrocinadora": desformatar_string_para_float(parcela_patrocinadora_str),
                            "Regime_de_tributacao": regime_tributacao,
                            "Recursos_portados": desformatar_string_para_float(recursos_portados_str),
                            "debito": desformatar_string_para_float(debito_str),
                            "Data_base_portabilidade": data_base_portabilidade
                        }
                        dados_completos_apos_salvar = salvar_dados_completos(row['Nome'], dados_atualizados)
                        st.success(f"Dados de Termo de Portabilidade para {row['Nome']} salvos!")

                        pdf_path_port, docx_path_port = gerar_documento_portabilidade(dados_completos_apos_salvar)
                        st.session_state.download_docx_portabilidade = docx_path_port
                        st.rerun()

                if st.session_state.download_docx_portabilidade:
                    if os.path.exists(st.session_state.download_docx_portabilidade):
                        with open(st.session_state.download_docx_portabilidade, "rb") as file:
                            st.download_button(
                                label="📥 Download Termo de Portabilidade DOCX",
                                data=file,
                                file_name=os.path.basename(st.session_state.download_docx_portabilidade),
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                key=f"download_btn_termo_port_{idx}"
                            )
                    else:
                        st.error(f"Arquivo DOCX não encontrado em: {st.session_state.download_docx_portabilidade}")

with tab6: # Carta de Portabilidade
    st.header("📧 Carta de Portabilidade")
    df_carta_portabilidade = carregar_dados()
    df_carta = df_carta_portabilidade[df_carta_portabilidade["Etapa"] == "Carta de Portabilidade"]

    if df_carta.empty:
        st.info("Nenhum formulário na etapa 'Carta de Portabilidade'.")
    else:
        for idx, row in df_carta.iterrows():
            with st.expander(f"Detalhes de {row['Nome']} - Matrícula: {formatar_matricula(row['Matricula'])}", expanded=False):
                with st.form(f"form_carta_portabilidade_{row['Nome']}"):
                    st.write(f"Preenchendo dados para Carta de Portabilidade para **{row['Nome']}** (Matrícula: {formatar_matricula(row['Matricula'])})")

                    data_transferencia_carta = st.text_input("Data de Transferência (dd/mm/aaaa)", value=row.get('Data_de_Transferencia_Carta', ''), key=f"data_transf_c_{idx}")
                    banco_carta = st.text_input("Banco (para carta)", value=row.get('Banco_Carta', ''), key=f"banco_c_{idx}")
                    agencia_carta = st.text_input("Agência (para carta)", value=row.get('Agencia_Carta', ''), key=f"ag_c_{idx}")
                    conta_corrente_carta = st.text_input("Conta Corrente (para carta)", value=row.get('Conta_Corrente_Carta', ''), key=f"cc_c_{idx}")
                    
                    st.subheader("Dados do Participante (revisar para a Carta)")
                    st.text_input("Nome do Participante", value=row.get('Nome', ''), disabled=True)
                    st.text_input("CPF do Participante", value=row.get('CPF', ''), disabled=True)
                    st.text_input("Matrícula do Participante", value=formatar_matricula(row.get('Matricula', '')), disabled=True)
                    st.text_input("Plano Original", value=row.get('Plano', ''), disabled=True)
                    
                    st.text_input("Rua", value=row.get('Rua', ''), key=f"rua_carta_{idx}")
                    # CORREÇÃO AQUI: Garante que o valor é string antes de chamar replace
                    complemento = st.text_input("Complemento", value=str(row.get('Complemento', '')).replace('nan', ''), key=f"comp_carta_{idx}")
                    st.text_input("Bairro", value=row.get('Bairro', ''), key=f"bairro_carta_{idx}")
                    st.text_input("CEP", value=row.get('CEP', ''), key=f"cep_carta_{idx}")
                    st.text_input("Cidade", value=row.get('Cidade', ''), key=f"cidade_carta_{idx}")
                    st.text_input("UF", value=row.get('UF', ''), key=f"uf_carta_{idx}")
                    st.text_input("Nº Ref. Documento", value=row.get('NRefDoc', ''), key=f"nref_carta_{idx}")

                    submitted_carta_port = st.form_submit_button("💾 Salvar Dados e Gerar DOCX (Carta de Portabilidade)")
                    if submitted_carta_port:
                        dados_atualizados = {
                            "Data_de_Transferencia_Carta": data_transferencia_carta,
                            "Banco_Carta": banco_carta,
                            "Agencia_Carta": agencia_carta,
                            "Conta_Corrente_Carta": conta_corrente_carta,
                            "Rua": st.session_state[f"rua_carta_{idx}"], "Complemento": st.session_state[f"comp_carta_{idx}"],
                            "Bairro": st.session_state[f"bairro_carta_{idx}"], "CEP": st.session_state[f"cep_carta_{idx}"],
                            "Cidade": st.session_state[f"cidade_carta_{idx}"], "UF": st.session_state[f"uf_carta_{idx}"],
                            "NRefDoc": st.session_state[f"nref_carta_{idx}"]
                        }
                        dados_completos_apos_salvar = salvar_dados_completos(row['Nome'], dados_atualizados)
                        st.success(f"Dados de Carta de Portabilidade para {row['Nome']} salvos!")
                        
                        pdf_path_carta, docx_path_carta = gerar_documento_carta_portabilidade(dados_completos_apos_salvar)
                        st.session_state.download_docx_carta = docx_path_carta
                        st.rerun()
                
                if st.session_state.download_docx_carta:
                    if os.path.exists(st.session_state.download_docx_carta):
                        with open(st.session_state.download_docx_carta, "rb") as file:
                            st.download_button(
                                label="📥 Download Carta de Portabilidade DOCX",
                                data=file,
                                file_name=os.path.basename(st.session_state.download_docx_carta),
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                key=f"download_btn_carta_{idx}"
                            )
                    else:
                        st.error(f"Arquivo DOCX não encontrado em: {st.session_state.download_docx_carta}")